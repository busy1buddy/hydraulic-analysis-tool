"""
DOCX Report Generator for EPANET Hydraulic Analysis
=====================================================
Generates professional Word documents with network summaries,
steady-state results, compliance checks, and transient analysis.
"""

from datetime import date

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Styling helpers
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn('w:shd'),
        {qn('w:fill'): color_hex, qn('w:val'): 'clear'},
    )
    shading.append(shading_elem)


def _shade_header_row(table, color_hex='2E4057'):
    """Apply dark shading to the first row of a table and set white text."""
    for cell in table.rows[0].cells:
        _set_cell_shading(cell, color_hex)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True


def _add_styled_table(doc, headers, rows, col_widths=None):
    """Create a table with styled header row and return it."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, value in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    _shade_header_row(table)
    return table


# ---------------------------------------------------------------------------
# Main report generator
# ---------------------------------------------------------------------------

def generate_docx_report(results, network_summary, output_path,
                         title='Hydraulic Analysis Report',
                         engineer_name='', project_name=''):
    """
    Generate a professional Word document report.

    Parameters
    ----------
    results : dict
        Combined analysis results. Expected keys (all optional):
        - steady_state : dict with pressures, flows, compliance
        - transient : dict with junctions, compliance, mitigation
        - fire_flow : dict with residual_pressures, compliance
        - water_quality : dict with junction_quality, compliance
    network_summary : dict
        Network metadata including node/link details. Expected keys:
        - junctions, reservoirs, tanks, pipes, valves (counts)
        - nodes : list of dicts with id, type, elevation, demand_lps / head
        - links : list of dicts with id, type, start, end, length,
          diameter_mm, roughness
    output_path : str
        Full path for the generated .docx file.
    title : str
        Report title shown on cover page.
    engineer_name : str
        Engineer name for the cover page.
    project_name : str
        Project name; overrides *title* on the cover page when provided.
    """
    doc = Document()

    # -- Default font --
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # =====================================================================
    # COVER PAGE
    # =====================================================================
    for _ in range(6):
        doc.add_paragraph('')

    cover_title = project_name if project_name else title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(cover_title)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Prepared using EPANET Hydraulic Analysis Toolkit')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)

    doc.add_paragraph('')

    if engineer_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'Engineer: {engineer_name}')
        run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Date: {date.today().strftime("%d %B %Y")}')
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Analysis per WSAA Guidelines')
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)

    doc.add_page_break()

    # =====================================================================
    # SECTION 1: NETWORK DESCRIPTION
    # =====================================================================
    doc.add_heading('1. Network Description', level=1)

    doc.add_heading('1.1 Network Summary', level=2)
    summary_rows = [
        ['Junctions', str(network_summary.get('junctions', 0))],
        ['Reservoirs', str(network_summary.get('reservoirs', 0))],
        ['Tanks', str(network_summary.get('tanks', 0))],
        ['Pipes', str(network_summary.get('pipes', 0))],
        ['Valves', str(network_summary.get('valves', 0))],
        ['Pumps', str(network_summary.get('pumps', 0))],
        ['Duration (hrs)', str(network_summary.get('duration_hrs', '-'))],
    ]
    _add_styled_table(doc, ['Component', 'Count'], summary_rows)

    # -- Node table --
    nodes = network_summary.get('nodes', [])
    if nodes:
        doc.add_heading('1.2 Node Details', level=2)
        node_rows = []
        for n in nodes:
            ntype = n.get('type', '')
            elevation = n.get('elevation', n.get('head', '-'))
            demand = n.get('demand_lps', '-')
            node_rows.append([n['id'], ntype.capitalize(), str(elevation), str(demand)])
        _add_styled_table(doc, ['ID', 'Type', 'Elevation (m)', 'Demand (LPS)'], node_rows)

    # -- Link table --
    links = network_summary.get('links', [])
    if links:
        doc.add_heading('1.3 Pipe / Link Details', level=2)
        link_rows = []
        for lk in links:
            link_rows.append([
                lk['id'],
                lk.get('start', '-'),
                lk.get('end', '-'),
                str(lk.get('length', '-')),
                str(lk.get('diameter_mm', '-')),
                str(lk.get('roughness', '-')),
            ])
        _add_styled_table(
            doc,
            ['ID', 'Start', 'End', 'Length (m)', 'Diameter (mm)', 'Roughness'],
            link_rows,
        )

    # =====================================================================
    # SECTION 2: STEADY-STATE RESULTS
    # =====================================================================
    steady = results.get('steady_state')
    if steady:
        doc.add_page_break()
        doc.add_heading('2. Steady-State Results', level=1)

        # Pressure table
        pressures = steady.get('pressures', {})
        if pressures:
            doc.add_heading('2.1 Junction Pressures', level=2)
            p_rows = []
            for junc, data in pressures.items():
                p_rows.append([
                    junc,
                    str(data.get('min_m', '-')),
                    str(data.get('max_m', '-')),
                    str(data.get('avg_m', '-')),
                ])
            _add_styled_table(doc, ['Junction', 'Min (m)', 'Max (m)', 'Avg (m)'], p_rows)

        # Flow table
        flows = steady.get('flows', {})
        if flows:
            doc.add_heading('2.2 Pipe Flows', level=2)
            f_rows = []
            for pipe, data in flows.items():
                f_rows.append([
                    pipe,
                    str(data.get('min_lps', '-')),
                    str(data.get('max_lps', '-')),
                    str(data.get('avg_lps', '-')),
                    str(data.get('avg_velocity_ms', '-')),
                ])
            _add_styled_table(
                doc,
                ['Pipe', 'Min (LPS)', 'Max (LPS)', 'Avg (LPS)', 'Velocity (m/s)'],
                f_rows,
            )

        # Summary paragraph
        n_junctions = len(pressures)
        n_pipes = len(flows)
        doc.add_paragraph(
            f'Steady-state analysis completed for {n_junctions} junctions '
            f'and {n_pipes} pipes. See compliance section for detailed checks.'
        )

    # =====================================================================
    # SECTION 3: COMPLIANCE
    # =====================================================================
    all_compliance = _collect_compliance(results)
    if all_compliance:
        doc.add_page_break()
        doc.add_heading('3. Compliance Summary', level=1)
        doc.add_paragraph(
            'The following compliance checks have been performed against '
            'Australian standards (WSAA Guidelines).'
        )

        comp_rows = []
        for item in all_compliance:
            ctype = item.get('type', 'INFO')
            element = item.get('element', '-')
            message = item.get('message', '')
            comp_rows.append([ctype, element, message])

        table = doc.add_table(rows=1 + len(comp_rows), cols=3)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, h in enumerate(['Status', 'Element', 'Message']):
            cell = table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shade_header_row(table)

        for r_idx, row_data in enumerate(comp_rows):
            for c_idx, value in enumerate(row_data):
                cell = table.rows[r_idx + 1].cells[c_idx]
                p = cell.paragraphs[0]
                run = p.add_run(str(value))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Colour-code status column
                if c_idx == 0:
                    run.bold = True
                    if value == 'CRITICAL':
                        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                    elif value == 'WARNING':
                        run.font.color.rgb = RGBColor(0xCC, 0x7A, 0x00)
                    elif value == 'OK':
                        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    # =====================================================================
    # SECTION 4: TRANSIENT RESULTS
    # =====================================================================
    transient = results.get('transient')
    if transient:
        doc.add_page_break()
        doc.add_heading('4. Transient Analysis (Water Hammer)', level=1)

        doc.add_paragraph(
            f'Valve: {transient.get("valve", "-")}  |  '
            f'Closure time: {transient.get("closure_time_s", "-")} s  |  '
            f'Wave speed: {transient.get("wave_speed_ms", "-")} m/s'
        )

        # Surge summary
        doc.add_heading('4.1 Surge Summary', level=2)
        doc.add_paragraph(
            f'Maximum surge: {transient.get("max_surge_m", "-")} m '
            f'({transient.get("max_surge_kPa", "-")} kPa)'
        )

        # Junction table
        junctions = transient.get('junctions', {})
        if junctions:
            doc.add_heading('4.2 Junction Surge Data', level=2)
            t_rows = []
            for junc, data in junctions.items():
                t_rows.append([
                    junc,
                    str(data.get('steady_head_m', '-')),
                    str(data.get('max_head_m', '-')),
                    str(data.get('min_head_m', '-')),
                    str(data.get('surge_m', '-')),
                    str(data.get('surge_kPa', '-')),
                ])
            _add_styled_table(
                doc,
                ['Junction', 'Steady (m)', 'Max (m)', 'Min (m)',
                 'Surge (m)', 'Surge (kPa)'],
                t_rows,
            )

        # Mitigation
        mitigation = transient.get('mitigation', [])
        if mitigation:
            doc.add_heading('4.3 Mitigation Recommendations', level=2)
            for rec in mitigation:
                doc.add_paragraph(rec, style='List Bullet')

    # =====================================================================
    # SECTION 5: FIRE FLOW RESULTS
    # =====================================================================
    fire_flow = results.get('fire_flow')
    if fire_flow:
        doc.add_page_break()
        doc.add_heading('5. Fire Flow Analysis', level=1)

        doc.add_paragraph(
            f'Fire flow node: {fire_flow.get("fire_node", "-")}  |  '
            f'Applied flow: {fire_flow.get("fire_flow_lps", "-")} LPS  |  '
            f'Fire node pressure: {fire_flow.get("fire_node_pressure_m", "-")} m'
        )

        residuals = fire_flow.get('residual_pressures', {})
        if residuals:
            doc.add_heading('5.1 Residual Pressures', level=2)
            ff_rows = [[junc, str(p)] for junc, p in residuals.items()]
            _add_styled_table(doc, ['Junction', 'Residual Pressure (m)'], ff_rows)

    # =====================================================================
    # SECTION 6: WATER QUALITY RESULTS
    # =====================================================================
    wq = results.get('water_quality')
    if wq:
        doc.add_page_break()
        doc.add_heading('6. Water Quality Analysis', level=1)

        doc.add_paragraph(
            f'Parameter: {wq.get("parameter", "-")}  |  '
            f'Simulation duration: {wq.get("duration_hrs", "-")} hrs'
        )

        jq = wq.get('junction_quality', {})
        if jq:
            doc.add_heading('6.1 Water Age at Junctions', level=2)
            wq_rows = []
            for junc, data in jq.items():
                wq_rows.append([
                    junc,
                    str(data.get('max_age_hrs', '-')),
                    str(data.get('avg_age_hrs', '-')),
                ])
            _add_styled_table(doc, ['Junction', 'Max Age (hrs)', 'Avg Age (hrs)'], wq_rows)

        stagnation = wq.get('stagnation_risk', [])
        if stagnation:
            doc.add_heading('6.2 Stagnation Risks', level=2)
            for junc in stagnation:
                doc.add_paragraph(junc, style='List Bullet')

    # =====================================================================
    # CONCLUSIONS
    # =====================================================================
    doc.add_page_break()
    doc.add_heading('Conclusions', level=1)

    conclusions = _build_conclusions(results)
    for para_text in conclusions:
        doc.add_paragraph(para_text)

    # =====================================================================
    # FOOTER (page numbers)
    # =====================================================================
    _add_page_numbers(doc)

    # Save
    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _collect_compliance(results):
    """Gather all compliance items from every analysis type."""
    items = []
    for key in ('steady_state', 'transient', 'fire_flow', 'water_quality'):
        sub = results.get(key)
        if sub and 'compliance' in sub:
            items.extend(sub['compliance'])
    return items


def _build_conclusions(results):
    """Auto-generate conclusion paragraphs from results."""
    paragraphs = []
    all_compliance = _collect_compliance(results)

    ok_count = sum(1 for c in all_compliance if c.get('type') == 'OK')
    warn_count = sum(1 for c in all_compliance if c.get('type') == 'WARNING')
    crit_count = sum(1 for c in all_compliance if c.get('type') == 'CRITICAL')

    if crit_count:
        paragraphs.append(
            f'CRITICAL: {crit_count} critical issue(s) were identified that '
            f'require immediate engineering attention before the network can be '
            f'considered compliant with WSAA guidelines.'
        )
    if warn_count:
        paragraphs.append(
            f'{warn_count} warning(s) were identified. These should be reviewed '
            f'and addressed where practicable to improve network performance.'
        )
    if ok_count and not crit_count and not warn_count:
        paragraphs.append(
            'All analyses passed compliance checks. The network meets '
            'WSAA guideline requirements for the scenarios tested.'
        )

    if results.get('transient'):
        surge = results['transient'].get('max_surge_m', 0)
        paragraphs.append(
            f'Transient analysis recorded a maximum surge of {surge} m. '
            f'Refer to Section 4 for mitigation recommendations.'
        )

    if not paragraphs:
        paragraphs.append(
            'Analysis complete. Review the detailed results in the '
            'preceding sections for engineering assessment.'
        )

    return paragraphs


def _add_page_numbers(doc):
    """Insert automatic page numbers in the document footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # PAGE field
        run = p.add_run()
        fld_char_begin = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
        run._element.append(fld_char_begin)

        run2 = p.add_run()
        instr = run2._element.makeelement(qn('w:instrText'), {})
        instr.text = ' PAGE '
        run2._element.append(instr)

        run3 = p.add_run()
        fld_char_end = run3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
        run3._element.append(fld_char_end)
