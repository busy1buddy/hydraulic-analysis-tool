"""
Tests for EPANET report generation (DOCX and PDF).
"""

import os
import sys
import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


# =========================================================================
# DOCX TESTS
# =========================================================================

class TestDocxReportCreated:
    """Verify that a DOCX report file is generated and non-empty."""

    def test_docx_report_created(self, loaded_network, steady_results):
        """Generate a DOCX report and verify the file exists and has content."""
        report_path = loaded_network.generate_report(
            format='docx',
            steady_results=steady_results,
            engineer_name='Test Engineer',
            project_name='Test Project',
        )
        assert os.path.isfile(report_path), f'Report not created at {report_path}'
        assert os.path.getsize(report_path) > 0, 'Report file is empty'

    def test_docx_with_no_results(self, loaded_network):
        """DOCX report with no analysis results should still generate."""
        report_path = loaded_network.generate_report(
            format='docx',
            engineer_name='Engineer',
            project_name='Empty Report',
        )
        assert os.path.isfile(report_path)
        assert os.path.getsize(report_path) > 0


class TestReportContainsSections:
    """Open the DOCX and verify it contains expected headings."""

    def test_report_contains_sections(self, loaded_network, steady_results):
        """Verify that key section headings are present in the report."""
        from docx import Document

        report_path = loaded_network.generate_report(
            format='docx',
            steady_results=steady_results,
            engineer_name='J. Smith',
            project_name='Section Test',
        )
        doc = Document(report_path)

        # Collect all heading text
        headings = []
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                headings.append(para.text)

        # Check for key sections
        heading_text = ' '.join(headings).lower()
        assert 'network description' in heading_text
        assert 'steady-state' in heading_text
        assert 'compliance' in heading_text
        assert 'conclusions' in heading_text

    def test_report_has_tables(self, loaded_network, steady_results):
        """Verify the DOCX report contains at least one table."""
        from docx import Document

        report_path = loaded_network.generate_report(
            format='docx',
            steady_results=steady_results,
        )
        doc = Document(report_path)
        assert len(doc.tables) >= 1, 'Report should contain at least one table'

    def test_report_cover_page_content(self, loaded_network, steady_results):
        """Verify cover page has project name and engineer."""
        from docx import Document

        report_path = loaded_network.generate_report(
            format='docx',
            steady_results=steady_results,
            engineer_name='Alice B.',
            project_name='Cover Page Test',
        )
        doc = Document(report_path)

        full_text = ' '.join(p.text for p in doc.paragraphs)
        assert 'Cover Page Test' in full_text
        assert 'Alice B.' in full_text
        assert 'WSAA' in full_text


class TestApiGenerateReport:
    """Test the HydraulicAPI.generate_report() method."""

    def test_api_generate_report(self, loaded_network):
        """Run steady-state, then generate a DOCX via the API method."""
        results = loaded_network.run_steady_state(save_plot=False)
        path = loaded_network.generate_report(
            format='docx',
            steady_results=results,
            project_name='API Test Report',
        )
        assert isinstance(path, str)
        assert path.endswith('.docx')
        assert os.path.isfile(path)

    def test_api_generate_pdf_report(self, loaded_network):
        """Run steady-state, then generate a PDF via the API method."""
        results = loaded_network.run_steady_state(save_plot=False)
        path = loaded_network.generate_report(
            format='pdf',
            steady_results=results,
            project_name='PDF Test Report',
        )
        assert isinstance(path, str)
        # Could be .pdf or .html depending on fpdf2 availability
        assert os.path.isfile(path)
        assert os.path.getsize(path) > 0

    def test_api_unsupported_format(self, loaded_network):
        """Unsupported format should return an error dict."""
        result = loaded_network.generate_report(format='xml')
        assert isinstance(result, dict)
        assert 'error' in result

    def test_api_no_network_loaded(self, api_instance):
        """generate_report without a loaded network returns error."""
        result = api_instance.generate_report(format='docx')
        assert isinstance(result, dict)
        assert 'error' in result

    def test_api_report_with_transient(self, transient_network, transient_results):
        """Generate report that includes transient results."""
        steady = transient_network.run_steady_state(save_plot=False)
        path = transient_network.generate_report(
            format='docx',
            steady_results=steady,
            transient_results=transient_results,
            project_name='Transient Report',
        )
        assert os.path.isfile(path)

        from docx import Document
        doc = Document(path)
        headings = [p.text.lower() for p in doc.paragraphs
                     if p.style.name.startswith('Heading')]
        heading_text = ' '.join(headings)
        assert 'transient' in heading_text
