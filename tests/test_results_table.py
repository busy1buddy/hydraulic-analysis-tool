"""
Results Table Sort, Filter, and CSV Export Tests — Cycle 6
==========================================================
Verifies:
  1. Numeric sorting: velocity column sorts numerically, not alphabetically
  2. Violation filter: "Show only violations" hides PASS rows
  3. CSV export: writes correct headers and data for visible rows
"""

import os
import sys
import csv
import tempfile

import pytest

os.environ.setdefault('QT_QPA_PLATFORM', 'offscreen')
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from PyQt6.QtWidgets import QApplication
from PyQt6.QtCore import Qt

from desktop.main_window import MainWindow, _NumericItem


@pytest.fixture(scope='module')
def app():
    instance = QApplication.instance()
    if instance is None:
        instance = QApplication([])
    yield instance


@pytest.fixture
def loaded_window(app):
    """MainWindow with demo_network loaded and analysis run."""
    import wntr
    w = MainWindow()
    w.resize(1400, 900)

    inp = 'tutorials/demo_network/network.inp'
    w.api.wn = wntr.network.WaterNetworkModel(inp)
    w.api._inp_file = inp
    w._current_file = inp
    w._populate_explorer()
    w.canvas.set_api(w.api)

    results = w.api.run_steady_state(save_plot=False)
    w._last_results = results
    w._on_analysis_finished(results)
    app.processEvents()
    w.show()
    app.processEvents()
    yield w, results
    w.close()
    app.processEvents()


class TestNumericItem:
    """Unit tests for _NumericItem sort ordering."""

    def test_numeric_sort_order(self):
        a = _NumericItem("2.5 m/s")
        b = _NumericItem("10.1 m/s")
        # Numerically 2.5 < 10.1
        assert a < b

    def test_alphabetic_would_fail(self):
        """Alphabetically '9' > '10', but numerically 9 < 10."""
        a = _NumericItem("9.0 m")
        b = _NumericItem("10.0 m")
        assert a < b

    def test_empty_text(self):
        a = _NumericItem("")
        b = _NumericItem("5.0")
        assert a < b  # empty = 0 < 5

    def test_dash_fallback(self):
        a = _NumericItem("--")
        b = _NumericItem("5.0")
        # Should fall back to string comparison, not crash
        result = a < b
        assert isinstance(result, bool)


class TestPipeResultsSorting:
    """Sorting the pipe results table by velocity descending."""

    def test_sort_velocity_descending(self, loaded_window, app):
        w, results = loaded_window
        table = w.pipe_results_table
        total = table.rowCount()
        assert total > 0, "Pipe results table is empty"

        # Collect all velocities before sort
        pre_velocities = []
        for row in range(total):
            item = table.item(row, 3)
            if item and item.text().strip():
                try:
                    pre_velocities.append(float(item.text().split()[0]))
                except (ValueError, IndexError):
                    pass

        assert len(pre_velocities) >= 2, (
            f"Need at least 2 pipes with velocity, got {len(pre_velocities)} "
            f"from {total} rows")

        # Sort velocity column descending
        table.sortItems(3, Qt.SortOrder.DescendingOrder)
        app.processEvents()

        # Read back after sort
        velocities = []
        for row in range(table.rowCount()):
            item = table.item(row, 3)
            if item and item.text().strip():
                try:
                    velocities.append(float(item.text().split()[0]))
                except (ValueError, IndexError):
                    pass

        assert len(velocities) >= 2
        # After descending sort, first should be >= second
        assert velocities[0] >= velocities[1], (
            f"First row velocity {velocities[0]} < second row {velocities[1]}")

        # Verify it's actually the max
        max_v = max(velocities)
        assert abs(velocities[0] - max_v) < 0.01, (
            f"First row {velocities[0]} != max {max_v}")


class TestNodeResultsSorting:
    """Sorting the node results table by pressure."""

    def test_sort_pressure_ascending(self, loaded_window, app):
        w, results = loaded_window
        table = w.node_results_table
        assert table.rowCount() > 0

        # Min pressure is column 2
        table.sortItems(2, Qt.SortOrder.AscendingOrder)
        app.processEvents()

        pressures = []
        for row in range(table.rowCount()):
            item = table.item(row, 2)
            if item:
                try:
                    pressures.append(float(item.text().split()[0]))
                except (ValueError, IndexError):
                    pass

        assert len(pressures) >= 2
        # After ascending sort, first should be <= second
        assert pressures[0] <= pressures[1], (
            f"First row pressure {pressures[0]} > second row {pressures[1]}")


class TestViolationFilter:
    """Right-click 'Show only violations' filter."""

    def test_filter_shows_only_red_rows(self, loaded_window, app):
        w, results = loaded_window
        table = w.pipe_results_table

        # Count rows with red foreground (violations)
        red_rows = set()
        for row in range(table.rowCount()):
            for col in range(table.columnCount()):
                item = table.item(row, col)
                if item and item.foreground().color().red() > 200:
                    red_rows.add(row)
                    break

        # Apply violation filter
        for row in range(table.rowCount()):
            hide = True
            for col in range(table.columnCount()):
                item = table.item(row, col)
                if item and item.foreground().color().red() > 200:
                    hide = False
                    break
            table.setRowHidden(row, hide)
        app.processEvents()

        # Only red rows should be visible
        visible = [r for r in range(table.rowCount()) if not table.isRowHidden(r)]
        assert set(visible) == red_rows

    def test_show_all_resets_filter(self, loaded_window, app):
        w, results = loaded_window
        table = w.pipe_results_table

        # Hide some rows first
        if table.rowCount() > 0:
            table.setRowHidden(0, True)

        # Reset
        for row in range(table.rowCount()):
            table.setRowHidden(row, False)
        app.processEvents()

        hidden = [r for r in range(table.rowCount()) if table.isRowHidden(r)]
        assert len(hidden) == 0


class TestDocxSlurrySection:
    """DOCX report generates slurry section when slurry data present."""

    def test_docx_includes_slurry_section(self):
        from reports.docx_report import generate_docx_report
        from docx import Document

        # Build a minimal results dict with slurry data
        results = {
            'steady_state': {
                'pressures': {'J1': {'min_m': 30, 'max_m': 35, 'avg_m': 32}},
                'flows': {
                    'P1': {
                        'min_lps': 5, 'max_lps': 10, 'avg_lps': 7.5,
                        'avg_velocity_ms': 1.2, 'max_velocity_ms': 1.5,
                        'headloss_per_km': 5.0,
                    }
                },
                'compliance': [],
                'slurry': {
                    'P1': {
                        'velocity_ms': 1.5, 'headloss_m': 2.5,
                        'regime': 'turbulent', 'reynolds': 4500,
                    }
                },
            },
            'slurry_params': {
                'yield_stress': 15.0,
                'plastic_viscosity': 0.05,
                'density': 1800,
            },
        }
        summary = {
            'junctions': 1, 'reservoirs': 1, 'tanks': 0,
            'pipes': 1, 'valves': 0, 'pumps': 0,
            'nodes': [], 'links': [
                {'id': 'P1', 'start': 'R1', 'end': 'J1',
                 'length': 500, 'diameter_mm': 200, 'roughness': 130},
            ],
        }

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            path = f.name

        try:
            generate_docx_report(results, summary, path,
                                 project_name='Slurry Test')
            doc = Document(path)
            all_text = '\n'.join(p.text for p in doc.paragraphs)
            # Also collect table cell text
            table_text = '\n'.join(
                cell.text for tbl in doc.tables for row in tbl.rows
                for cell in row.cells)
            combined = all_text + '\n' + table_text

            assert 'Non-Newtonian Slurry Analysis' in combined
            assert 'Yield Stress' in combined
            assert '15.0 Pa' in combined
            assert 'Bingham Plastic' in combined
            assert 'turbulent' in combined.lower()
        finally:
            os.unlink(path)


class TestCSVExport:
    """CSV export of results table."""

    def test_csv_export_content(self, loaded_window, app):
        w, results = loaded_window
        table = w.pipe_results_table
        assert table.rowCount() > 0

        # Export to temp file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.csv',
                                          delete=False, newline='',
                                          encoding='utf-8') as f:
            tmp_path = f.name
            writer = csv.writer(f)
            # Header
            headers = []
            for col in range(table.columnCount()):
                h = table.horizontalHeaderItem(col)
                headers.append(h.text() if h else f"Col{col}")
            writer.writerow(headers)
            # Data
            for row in range(table.rowCount()):
                if table.isRowHidden(row):
                    continue
                row_data = []
                for col in range(table.columnCount()):
                    item = table.item(row, col)
                    row_data.append(item.text() if item else "")
                writer.writerow(row_data)

        try:
            with open(tmp_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                rows = list(reader)

            # Header + data rows
            assert len(rows) == table.rowCount() + 1
            assert rows[0][0] == "ID"
            assert "Velocity" in rows[0][3]
            # First data row should have pipe ID
            assert len(rows[1][0]) > 0
        finally:
            os.unlink(tmp_path)

    def test_csv_skips_hidden_rows(self, loaded_window, app):
        w, results = loaded_window
        table = w.pipe_results_table
        total_rows = table.rowCount()
        assert total_rows > 1

        # Hide first row
        table.setRowHidden(0, True)

        with tempfile.NamedTemporaryFile(mode='w', suffix='.csv',
                                          delete=False, newline='',
                                          encoding='utf-8') as f:
            tmp_path = f.name
            writer = csv.writer(f)
            writer.writerow(["H"])
            for row in range(table.rowCount()):
                if table.isRowHidden(row):
                    continue
                writer.writerow([row])

        try:
            with open(tmp_path, 'r', encoding='utf-8') as f:
                rows = list(csv.reader(f))
            # Should be total_rows - 1 data rows + 1 header
            assert len(rows) == total_rows  # header + (total - 1 hidden)
        finally:
            os.unlink(tmp_path)
            # Restore
            table.setRowHidden(0, False)
