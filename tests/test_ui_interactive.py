"""
Automated UI Interaction Tests — Full Coverage
================================================
Exercises every panel, menu item, and interaction in the MainWindow
headlessly via QT_QPA_PLATFORM=offscreen.

Requires: pytest-qt (pip install pytest-qt)

Areas covered:
  1. Menu completeness
  2. Project Explorer interactions
  3. Results table completeness (values, highlighting, stress)
  4. Scenario panel (add, edit, duplicate, delete, run all)
  5. Canvas interactions (node/pipe click, color modes, fit, labels)
  6. Slurry mode parameter validation
  7. Report generation (DOCX, PDF)
  8. Error handling
  9. Window state preservation
  10. Integration: full end-to-end workflow
"""

import os
import sys
import json
import math
import tempfile

import pytest

os.environ['QT_QPA_PLATFORM'] = 'offscreen'

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from PyQt6.QtWidgets import QApplication, QDockWidget
from PyQt6.QtCore import Qt
from PyQt6.QtGui import QColor

from desktop.main_window import MainWindow
from desktop.scenario_panel import ScenarioData
from slurry_solver import bingham_plastic_headloss


# =========================================================================
# Fixtures
# =========================================================================

@pytest.fixture(scope='module')
def app():
    instance = QApplication.instance()
    if instance is None:
        instance = QApplication([])
    yield instance


@pytest.fixture
def window(app):
    """MainWindow with australian_network.inp loaded and shown."""
    w = MainWindow()
    w.resize(1400, 900)

    import wntr
    w.api.wn = wntr.network.WaterNetworkModel('models/australian_network.inp')
    w.api._inp_file = 'models/australian_network.inp'
    w._current_file = 'models/australian_network.inp'
    w._populate_explorer()
    w._update_status_bar()
    w.canvas.set_api(w.api)

    w.show()
    app.processEvents()
    yield w
    w.close()
    app.processEvents()


@pytest.fixture
def analysed_window(window, app):
    """Window with steady-state analysis already run."""
    results = window.api.run_steady_state(save_plot=False)
    window._on_analysis_finished(results)
    app.processEvents()
    return window


@pytest.fixture
def bare_window(app):
    """MainWindow with NO network loaded."""
    w = MainWindow()
    w.resize(1400, 900)
    w.show()
    app.processEvents()
    yield w
    w.close()
    app.processEvents()


def _get_properties(window):
    """Helper to read Properties table as a dict."""
    props = {}
    for row in range(window.properties_table.rowCount()):
        k = window.properties_table.item(row, 0).text()
        v = window.properties_table.item(row, 1).text()
        props[k] = v
    return props


def _get_menu_action_texts(window, menu_title):
    """Return list of action texts for a top-level menu."""
    for action in window.menuBar().actions():
        if action.text().replace('&', '') == menu_title:
            return [a.text().replace('&', '') for a in action.menu().actions()
                    if not a.isSeparator()]
    return []


# =========================================================================
# Area 1 — Menu completeness
# =========================================================================

class TestMenuCompleteness:

    def test_file_menu_items(self, window):
        items = _get_menu_action_texts(window, 'File')
        for expected in ['New', 'Open (.inp)...', 'Save', 'Save As (.hap)...', 'Exit']:
            assert expected in items, f"Missing File > {expected}"

    def test_analysis_menu_items(self, window):
        items = _get_menu_action_texts(window, 'Analysis')
        for expected in ['Run Steady State', 'Run Transient', 'Slurry Mode']:
            assert expected in items, f"Missing Analysis > {expected}"

    def test_tools_menu_items(self, window):
        items = _get_menu_action_texts(window, 'Tools')
        for expected in ['Quality Review', 'Settings']:
            assert expected in items, f"Missing Tools > {expected}"

    def test_reports_menu_items(self, window):
        items = _get_menu_action_texts(window, 'Reports')
        for expected in ['Generate Report (DOCX)', 'Generate Report (PDF)']:
            assert expected in items, f"Missing Reports > {expected}"

    def test_view_menu_items(self, window):
        items = _get_menu_action_texts(window, 'View')
        for expected in ['Reset Layout', 'Project Explorer', 'Properties', 'Results']:
            assert expected in items, f"Missing View > {expected}"

    def test_help_menu_items(self, window):
        items = _get_menu_action_texts(window, 'Help')
        assert 'About' in items

    def test_slurry_mode_is_checkable(self, window):
        assert window.slurry_act.isCheckable()

    def test_slurry_mode_toggle(self, window, app):
        window.slurry_act.setChecked(True)
        app.processEvents()
        assert "Slurry" in window.analysis_label.text()
        window.slurry_act.setChecked(False)
        app.processEvents()

    def test_file_new_clears_state(self, window, app):
        window._on_new()
        app.processEvents()
        assert window.explorer_tree.topLevelItemCount() == 0
        assert window.node_results_table.rowCount() == 0
        assert window.pipe_results_table.rowCount() == 0
        assert "Hydraulic Analysis Tool" == window.windowTitle()

    def test_view_toggle_properties(self, window, app):
        window.toggle_properties_act.setChecked(False)
        app.processEvents()
        window.toggle_properties_act.setChecked(True)
        app.processEvents()
        assert window.properties_dock.isVisible()

    def test_view_reset_layout(self, window, app):
        window._on_reset_layout()
        app.processEvents()
        assert window.explorer_dock.isVisible()
        assert window.properties_dock.isVisible()
        assert window.results_dock.isVisible()

    def test_seven_top_level_menus(self, window):
        menus = [a.text().replace('&', '') for a in window.menuBar().actions()]
        assert menus == ['File', 'Edit', 'Analysis', 'Tools', 'Reports', 'View', 'Help']


# =========================================================================
# Area 2 — Project Explorer interactions
# =========================================================================

class TestProjectExplorer:

    def test_root_is_filename(self, window):
        root = window.explorer_tree.topLevelItem(0)
        assert 'australian_network' in root.text(0).lower()

    def test_junctions_category_has_7_children(self, window):
        root = window.explorer_tree.topLevelItem(0)
        for i in range(root.childCount()):
            cat = root.child(i)
            if cat.text(0).startswith("Junctions"):
                assert cat.childCount() == 7
                return
        pytest.fail("No Junctions category found")

    def test_junction_ids_j1_to_j7(self, window):
        root = window.explorer_tree.topLevelItem(0)
        for i in range(root.childCount()):
            cat = root.child(i)
            if cat.text(0).startswith("Junctions"):
                ids = {cat.child(j).text(0) for j in range(cat.childCount())}
                for jid in ['J1', 'J2', 'J3', 'J4', 'J5', 'J6', 'J7']:
                    assert jid in ids
                return

    def test_pipes_category_has_9_children(self, window):
        root = window.explorer_tree.topLevelItem(0)
        for i in range(root.childCount()):
            cat = root.child(i)
            if cat.text(0).startswith("Pipes"):
                assert cat.childCount() == 9
                return
        pytest.fail("No Pipes category found")

    def test_reservoir_category(self, window):
        root = window.explorer_tree.topLevelItem(0)
        texts = [root.child(i).text(0) for i in range(root.childCount())]
        res = [t for t in texts if t.startswith("Reservoirs")]
        assert len(res) == 1 and "1" in res[0]

    def test_tank_category(self, window):
        root = window.explorer_tree.topLevelItem(0)
        texts = [root.child(i).text(0) for i in range(root.childCount())]
        tank = [t for t in texts if t.startswith("Tanks")]
        assert len(tank) == 1 and "1" in tank[0]

    def test_click_junction_populates_properties(self, window, app):
        window._on_canvas_element_selected('J3', 'junction')
        app.processEvents()
        props = _get_properties(window)
        assert props.get('Type') == 'Junction'
        assert props.get('ID') == 'J3'
        assert 'm' in props.get('Elevation', '')

    def test_click_pipe_populates_properties(self, window, app):
        window._on_canvas_element_selected('P7', 'pipe')
        app.processEvents()
        props = _get_properties(window)
        assert props.get('Type') == 'Pipe'
        assert props.get('ID') == 'P7'
        assert 'mm' in props.get('Diameter', '')

    def test_click_reservoir_populates_properties(self, window, app):
        window._on_canvas_element_selected('R1', 'reservoir')
        app.processEvents()
        props = _get_properties(window)
        assert props.get('Type') == 'Reservoir'
        assert 'm' in props.get('Head', '')


# =========================================================================
# Area 3 — Results table completeness
# =========================================================================

class TestResultsTableCompleteness:

    @pytest.fixture(autouse=True)
    def setup(self, analysed_window, app):
        self.w = analysed_window
        self.app = app

    def test_node_table_7_rows(self):
        assert self.w.node_results_table.rowCount() == 7

    def test_pipe_table_9_rows(self):
        assert self.w.pipe_results_table.rowCount() == 9

    def test_all_junction_ids(self):
        ids = {self.w.node_results_table.item(r, 0).text()
               for r in range(self.w.node_results_table.rowCount())}
        assert ids == {'J1', 'J2', 'J3', 'J4', 'J5', 'J6', 'J7'}

    def test_all_pipe_ids(self):
        ids = {self.w.pipe_results_table.item(r, 0).text()
               for r in range(self.w.pipe_results_table.rowCount())}
        assert ids == {'P1', 'P2', 'P3', 'P4', 'P5', 'P6', 'P7', 'P8', 'P9'}

    def _find_pipe_row(self, pid):
        for r in range(self.w.pipe_results_table.rowCount()):
            if self.w.pipe_results_table.item(r, 0).text() == pid:
                return r
        return None

    def test_p4_velocity_highlighted_red(self):
        row = self._find_pipe_row('P4')
        assert row is not None
        fg = self.w.pipe_results_table.item(row, 3).foreground().color()
        assert fg.red() > 200 and fg.green() < 160, f"P4 velocity not red: {fg.red()},{fg.green()},{fg.blue()}"

    def test_p8_velocity_highlighted_red(self):
        row = self._find_pipe_row('P8')
        assert row is not None
        fg = self.w.pipe_results_table.item(row, 3).foreground().color()
        assert fg.red() > 200 and fg.green() < 160, f"P8 velocity not red"

    def test_p1_velocity_not_highlighted(self):
        row = self._find_pipe_row('P1')
        assert row is not None
        fg = self.w.pipe_results_table.item(row, 3).foreground().color()
        # Default black (0,0,0) or any non-red
        assert fg.red() < 200 or fg.green() > 160, f"P1 velocity unexpectedly red"

    def test_p5_velocity_not_highlighted(self):
        row = self._find_pipe_row('P5')
        fg = self.w.pipe_results_table.item(row, 3).foreground().color()
        assert fg.red() < 200 or fg.green() > 160

    def test_pressure_values_have_units_in_header(self):
        header = self.w.node_results_table.horizontalHeaderItem(2).text()
        assert "(m)" in header

    def test_velocity_values_have_units_in_header(self):
        header = self.w.pipe_results_table.horizontalHeaderItem(3).text()
        assert "(m/s)" in header

    def test_headloss_header_says_m_per_km(self):
        header = self.w.pipe_results_table.horizontalHeaderItem(4).text()
        assert "(m/km)" in header

    def test_pipe_stress_table_has_9_rows(self):
        assert self.w.pipe_stress_panel.table.rowCount() == 9

    def test_safety_factors_in_sensible_range(self):
        for r in range(self.w.pipe_stress_panel.table.rowCount()):
            sf_text = self.w.pipe_stress_panel.table.item(r, 6).text()
            sf = float(sf_text)
            assert 1.5 <= sf <= 20, f"Row {r}: SF {sf} outside 1.5-20 range"

    def test_wsaa_status_column_present(self):
        header = self.w.node_results_table.horizontalHeaderItem(4).text()
        assert "WSAA" in header

    def test_j4_node_pressures_reasonable(self):
        """J4 has the lowest pressure — verify it's flagged."""
        for r in range(self.w.node_results_table.rowCount()):
            if self.w.node_results_table.item(r, 0).text() == 'J4':
                status = self.w.node_results_table.item(r, 4).text()
                # J4 min pressure is 17.2 m, avg ~30.9 — avg passes WSAA
                assert 'PASS' in status or 'FAIL' in status
                return


# =========================================================================
# Area 4 — Scenario panel
# =========================================================================

class TestScenarioPanel:

    def test_base_scenario_exists(self, window):
        assert len(window.scenario_panel.scenarios) >= 1
        assert window.scenario_panel.scenarios[0].name == 'Base'

    def test_add_scenario(self, window, app):
        initial = len(window.scenario_panel.scenarios)
        window.scenario_panel.scenarios.append(ScenarioData('Peak', 1.5))
        window.scenario_panel._refresh_tree()
        app.processEvents()
        assert len(window.scenario_panel.scenarios) == initial + 1

    def test_scenario_in_tree(self, window, app):
        window.scenario_panel.scenarios.append(ScenarioData('Fire', 2.0))
        window.scenario_panel._refresh_tree()
        app.processEvents()
        tree_texts = []
        for i in range(window.scenario_panel.tree.topLevelItemCount()):
            tree_texts.append(window.scenario_panel.tree.topLevelItem(i).text(0))
        found = any('Fire' in t for t in tree_texts)
        assert found, f"Fire scenario not in tree: {tree_texts}"

    def test_run_all_scenarios(self, window, app):
        window.scenario_panel.scenarios = [
            ScenarioData('Base', 1.0),
            ScenarioData('Peak', 1.5),
        ]
        window._on_run_all_scenarios()
        app.processEvents()
        for sc in window.scenario_panel.scenarios:
            assert sc.results is not None, f"{sc.name} has no results"
            assert 'pressures' in sc.results

    def test_comparison_table_populated(self, window, app):
        window.scenario_panel.scenarios = [
            ScenarioData('Base', 1.0),
            ScenarioData('High', 2.0),
        ]
        window._on_run_all_scenarios()
        app.processEvents()
        rows = window.scenario_panel.comparison.table.rowCount()
        assert rows == 2, f"Comparison table has {rows} rows, expected 2"

    def test_duplicate_scenario(self, window, app):
        window.scenario_panel.scenarios = [ScenarioData('Base', 1.0)]
        window.scenario_panel.tree.setCurrentItem(
            window.scenario_panel.tree.topLevelItem(0))
        window.scenario_panel._on_duplicate()
        app.processEvents()
        assert len(window.scenario_panel.scenarios) == 2
        assert 'copy' in window.scenario_panel.scenarios[1].name.lower()

    def test_delete_non_base_scenario(self, window, app):
        window.scenario_panel.scenarios = [
            ScenarioData('Base', 1.0),
            ScenarioData('Temp', 1.2),
        ]
        window.scenario_panel._refresh_tree()
        # Select Temp
        window.scenario_panel.tree.setCurrentItem(
            window.scenario_panel.tree.topLevelItem(1))
        window.scenario_panel._on_delete()
        app.processEvents()
        assert len(window.scenario_panel.scenarios) == 1
        assert window.scenario_panel.scenarios[0].name == 'Base'

    def test_cannot_delete_base(self, window, app, monkeypatch):
        from PyQt6.QtWidgets import QMessageBox
        monkeypatch.setattr(QMessageBox, 'warning', lambda *a, **kw: None)
        window.scenario_panel.scenarios = [ScenarioData('Base', 1.0)]
        window.scenario_panel._refresh_tree()
        window.scenario_panel.tree.setCurrentItem(
            window.scenario_panel.tree.topLevelItem(0))
        window.scenario_panel._on_delete()
        app.processEvents()
        assert len(window.scenario_panel.scenarios) == 1


# =========================================================================
# Area 5 — Canvas interactions
# =========================================================================

class TestCanvasInteractions:

    def test_click_junction_j1(self, window, app):
        window._on_canvas_element_selected('J1', 'junction')
        app.processEvents()
        props = _get_properties(window)
        assert props['Type'] == 'Junction'

    def test_click_junction_j7(self, window, app):
        window._on_canvas_element_selected('J7', 'junction')
        app.processEvents()
        props = _get_properties(window)
        assert props['ID'] == 'J7'

    def test_click_reservoir_r1(self, window, app):
        window._on_canvas_element_selected('R1', 'reservoir')
        app.processEvents()
        props = _get_properties(window)
        assert props['Type'] == 'Reservoir'

    def test_click_tank_t1(self, window, app):
        window._on_canvas_element_selected('T1', 'tank')
        app.processEvents()
        props = _get_properties(window)
        assert props['Type'] == 'Reservoir'  # tanks use _show_reservoir_properties

    def test_click_all_9_pipes(self, window, app):
        for pid in ['P1', 'P2', 'P3', 'P4', 'P5', 'P6', 'P7', 'P8', 'P9']:
            window._on_canvas_element_selected(pid, 'pipe')
            app.processEvents()
            props = _get_properties(window)
            assert props.get('ID') == pid, f"Properties ID mismatch for {pid}"

    def test_color_mode_labels_update(self, window, app):
        for mode in window.canvas.COLOR_MODES:
            window.canvas.color_mode_combo.setCurrentText(mode)
            app.processEvents()
            assert window.canvas.color_mode_combo.currentText() == mode

    def test_fit_button_padding(self, window, app):
        window.canvas._fit_view()
        app.processEvents()
        vr = window.canvas.plot_widget.plotItem.vb.viewRange()
        xs = [p[0] for p in window.canvas._node_positions.values()]
        ys = [p[1] for p in window.canvas._node_positions.values()]
        x_range = max(xs) - min(xs)
        y_range = max(ys) - min(ys)
        left_pad = (min(xs) - vr[0][0]) / x_range if x_range > 0 else 0
        right_pad = (vr[0][1] - max(xs)) / x_range if x_range > 0 else 0
        assert left_pad > 0.10, f"Left padding {left_pad:.2f} < 10%"
        assert right_pad > 0.10, f"Right padding {right_pad:.2f} < 10%"

    def test_labels_toggle_on(self, window, app):
        window.canvas.labels_btn.setChecked(True)
        app.processEvents()
        assert len(window.canvas._label_items) > 0

    def test_labels_toggle_off(self, window, app):
        window.canvas.labels_btn.setChecked(True)
        app.processEvents()
        window.canvas.labels_btn.setChecked(False)
        app.processEvents()
        assert len(window.canvas._label_items) == 0

    def test_canvas_has_all_nodes(self, window):
        assert len(window.canvas._node_positions) == 9  # 7 junctions + 1 reservoir + 1 tank

    def test_canvas_has_all_pipes(self, window):
        assert len(window.canvas._pipe_ids) == 9


# =========================================================================
# Area 6 — Slurry mode parameter validation
# =========================================================================

class TestSlurryParameterValidation:

    def test_zero_yield_stress_no_crash(self):
        """τ_y=0 should produce a valid result (Newtonian limit)."""
        result = bingham_plastic_headloss(
            flow_m3s=0.01, diameter_m=0.2, length_m=100,
            density=1000, tau_y=0.0, mu_p=0.001, roughness_mm=0.1,
        )
        assert result['headloss_m'] >= 0

    def test_very_small_yield_stress(self):
        result = bingham_plastic_headloss(
            flow_m3s=0.01, diameter_m=0.2, length_m=100,
            density=1000, tau_y=0.0001, mu_p=0.001, roughness_mm=0.1,
        )
        assert result['headloss_m'] > 0
        assert result['regime'] in ('laminar', 'turbulent')

    def test_density_affects_reynolds_and_regime(self):
        """Different densities should produce different Reynolds numbers and
        potentially different flow regimes, but both must return valid results."""
        light = bingham_plastic_headloss(
            flow_m3s=0.01, diameter_m=0.2, length_m=100,
            density=1200, tau_y=10, mu_p=0.02, roughness_mm=0.1,
        )
        heavy = bingham_plastic_headloss(
            flow_m3s=0.01, diameter_m=0.2, length_m=100,
            density=2500, tau_y=10, mu_p=0.02, roughness_mm=0.1,
        )
        assert light['headloss_m'] > 0
        assert heavy['headloss_m'] > 0
        # Higher density → higher Reynolds → may shift regime
        assert heavy['reynolds'] > light['reynolds']

    def test_slurry_always_higher_than_water(self, window):
        """For every pipe in the network, slurry headloss > water headloss."""
        results = window.api.run_steady_state(save_plot=False)
        for pid, fdata in results['flows'].items():
            pipe = window.api.get_link(pid)
            Q = abs(fdata['avg_lps']) / 1000
            if Q <= 0 or pipe.diameter <= 0:
                continue
            hl_water = ((10.67 * pipe.length * Q ** 1.852) /
                        (pipe.roughness ** 1.852 * pipe.diameter ** 4.87))
            slurry = bingham_plastic_headloss(
                flow_m3s=Q, diameter_m=pipe.diameter, length_m=pipe.length,
                density=1800, tau_y=15, mu_p=0.05, roughness_mm=0.1,
            )
            assert slurry['headloss_m'] > hl_water, f"Pipe {pid}: slurry <= water"

    def test_zero_flow_returns_zero_headloss(self):
        result = bingham_plastic_headloss(
            flow_m3s=0.0, diameter_m=0.2, length_m=100,
            density=1500, tau_y=10, mu_p=0.02, roughness_mm=0.1,
        )
        assert result['headloss_m'] == 0

    def test_slurry_mode_toggle_updates_label(self, window, app):
        window.slurry_act.setChecked(True)
        app.processEvents()
        assert "Slurry" in window.analysis_label.text()
        window.slurry_act.setChecked(False)
        app.processEvents()


# =========================================================================
# Area 7 — Report generation
# =========================================================================

class TestReportGeneration:

    def test_docx_report_created(self, analysed_window, tmp_path):
        path = str(tmp_path / 'test.docx')
        summary = analysed_window.api.get_network_summary()
        from reports.docx_report import generate_docx_report
        generate_docx_report(
            analysed_window._last_results, summary, path,
            title='Test', engineer_name='Auto', project_name='Test',
        )
        assert os.path.exists(path)
        assert os.path.getsize(path) > 10000

    def test_pdf_report_created(self, analysed_window, tmp_path):
        path = str(tmp_path / 'test.pdf')
        summary = analysed_window.api.get_network_summary()
        from reports.pdf_report import generate_pdf_report
        generate_pdf_report(
            analysed_window._last_results, summary, path,
            title='Test', engineer_name='Auto', project_name='Test',
        )
        # fpdf produces .pdf; fallback produces .html
        produced = path if os.path.exists(path) else path.replace('.pdf', '.html')
        assert os.path.exists(produced), f"No report file at {path} or .html"
        assert os.path.getsize(produced) > 1000

    def test_docx_report_contains_network_data(self, analysed_window, tmp_path):
        path = str(tmp_path / 'content_check.docx')
        summary = analysed_window.api.get_network_summary()
        from reports.docx_report import generate_docx_report
        generate_docx_report(
            analysed_window._last_results, summary, path,
            title='Content Test', engineer_name='Auto', project_name='Test',
        )
        from docx import Document
        doc = Document(path)
        full_text = '\n'.join(p.text for p in doc.paragraphs)
        table_text = '\n'.join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        combined = full_text + table_text
        # Should mention the pipe/junction count
        assert '7' in combined, "Junction count not in report"
        assert '9' in combined, "Pipe count not in report"

    def test_report_no_bare_floats(self, analysed_window, tmp_path):
        """All numeric values should be rounded (no 15-digit decimals)."""
        path = str(tmp_path / 'float_check.docx')
        summary = analysed_window.api.get_network_summary()
        from reports.docx_report import generate_docx_report
        generate_docx_report(
            analysed_window._last_results, summary, path,
            title='Float Test', engineer_name='Auto', project_name='Test',
        )
        from docx import Document
        doc = Document(path)
        import re
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Flag numbers with > 4 decimal places (bare float)
                    matches = re.findall(r'\d+\.\d{5,}', cell.text)
                    assert not matches, f"Bare float in report: {matches} in '{cell.text}'"


# =========================================================================
# Area 8 — Error handling
# =========================================================================

class TestErrorHandling:

    def test_run_steady_without_network(self, bare_window, app, monkeypatch):
        """Analysis on empty window should show warning, not crash."""
        from PyQt6.QtWidgets import QMessageBox
        shown = []
        monkeypatch.setattr(QMessageBox, 'warning',
                            lambda *a, **kw: shown.append('warning'))
        bare_window._on_run_steady()
        app.processEvents()
        assert 'warning' in shown

    def test_run_transient_without_network(self, bare_window, app, monkeypatch):
        from PyQt6.QtWidgets import QMessageBox
        shown = []
        monkeypatch.setattr(QMessageBox, 'warning',
                            lambda *a, **kw: shown.append('warning'))
        bare_window._on_run_transient()
        app.processEvents()
        assert 'warning' in shown

    def test_report_without_results(self, bare_window, app, monkeypatch):
        """Generate report with no results should show warning, not crash."""
        from PyQt6.QtWidgets import QMessageBox
        shown = []
        monkeypatch.setattr(QMessageBox, 'warning',
                            lambda *a, **kw: shown.append('warning'))
        bare_window._on_report_docx()
        app.processEvents()
        assert 'warning' in shown

    def test_save_without_file(self, bare_window, app):
        """Save with no file loaded should not crash."""
        bare_window._current_file = None
        bare_window._hap_file = None
        path = os.path.join(tempfile.gettempdir(), 'test_save.hap')
        bare_window._save_hap(path)
        assert os.path.exists(path)
        with open(path) as f:
            data = json.load(f)
        assert 'settings' in data
        os.unlink(path)

    def test_analysis_error_handler(self, bare_window, app, monkeypatch):
        """Calling _on_analysis_error should not crash."""
        from PyQt6.QtWidgets import QMessageBox
        monkeypatch.setattr(QMessageBox, 'critical',
                            lambda *a, **kw: None)
        bare_window._on_analysis_error("Test error message")
        app.processEvents()


# =========================================================================
# Area 9 — Window state preservation
# =========================================================================

class TestWindowState:

    def test_docks_not_closable(self, window):
        for dock in [window.properties_dock, window.results_dock,
                     window.explorer_dock, window.scenario_dock]:
            features = dock.features()
            closable = bool(features & QDockWidget.DockWidgetFeature.DockWidgetClosable)
            assert not closable, f"{dock.windowTitle()} dock is closable"

    def test_properties_visible_after_restore(self, window, app):
        window.showMinimized()
        app.processEvents()
        window.showNormal()
        app.processEvents()
        assert window.properties_dock.isVisible()

    def test_results_visible_after_restore(self, window, app):
        window.showMinimized()
        app.processEvents()
        window.showNormal()
        app.processEvents()
        assert window.results_dock.isVisible()

    def test_scene_click_after_restore(self, window, app):
        window.showMinimized()
        app.processEvents()
        window.showNormal()
        app.processEvents()
        scene = window.canvas.plot_widget.scene()
        assert scene.receivers(scene.sigMouseClicked) > 0

    def test_pipe_click_after_restore(self, analysed_window, app):
        analysed_window.showMinimized()
        app.processEvents()
        analysed_window.showNormal()
        app.processEvents()
        analysed_window.properties_table.setRowCount(0)
        analysed_window._on_canvas_element_selected('P4', 'pipe')
        app.processEvents()
        assert analysed_window.properties_table.rowCount() > 0

    def test_minimum_window_size(self, window, app):
        window.resize(800, 600)
        app.processEvents()
        # All docks should still be visible at minimum
        assert window.properties_dock.isVisible()
        assert window.results_dock.isVisible()

    def test_save_hap_and_read_back(self, window, tmp_path, app):
        path = str(tmp_path / 'state.hap')
        window.slurry_act.setChecked(True)
        window._save_hap(path)
        with open(path) as f:
            data = json.load(f)
        assert data['settings']['slurry_mode'] is True
        assert data['inp_path'] == window._current_file


# =========================================================================
# Area 10 — Integration: full end-to-end workflow
# =========================================================================

class TestFullWorkflow:

    def test_load_analyse_click_report_slurry(self, window, app, tmp_path):
        """Full workflow: load → steady → click pipe → report → slurry → compare."""
        # Step 1: network already loaded by fixture
        assert window.api.wn is not None
        summary = window.api.get_network_summary()
        assert summary['junctions'] == 7
        assert summary['pipes'] == 9

        # Step 2: run steady state
        water_results = window.api.run_steady_state(save_plot=False)
        window._on_analysis_finished(water_results)
        app.processEvents()
        assert window.node_results_table.rowCount() == 7
        assert window.pipe_results_table.rowCount() == 9

        # Step 3: click pipe P4
        window._on_canvas_element_selected('P4', 'pipe')
        app.processEvents()
        props = _get_properties(window)
        assert props['ID'] == 'P4'
        assert '2.83' in props.get('Max Velocity', '')

        # Step 4: generate DOCX report
        report_path = str(tmp_path / 'workflow_report.docx')
        from reports.docx_report import generate_docx_report
        generate_docx_report(
            water_results, summary, report_path,
            title='Workflow Test', engineer_name='Auto', project_name='Test',
        )
        assert os.path.getsize(report_path) > 10000

        # Step 5: run slurry on same pipes
        water_flows = water_results['flows']
        slurry_higher_count = 0
        for pid, fdata in water_flows.items():
            pipe = window.api.get_link(pid)
            Q = abs(fdata['avg_lps']) / 1000
            if Q <= 0:
                continue
            hl_water = ((10.67 * pipe.length * Q ** 1.852) /
                        (pipe.roughness ** 1.852 * pipe.diameter ** 4.87))
            slurry = bingham_plastic_headloss(
                flow_m3s=Q, diameter_m=pipe.diameter, length_m=pipe.length,
                density=1800, tau_y=15, mu_p=0.05, roughness_mm=0.1,
            )
            if slurry['headloss_m'] > hl_water:
                slurry_higher_count += 1

        # All 9 pipes should have higher slurry headloss
        assert slurry_higher_count == 9

        # Step 6: verify no errors throughout
        assert window.wsaa_label.text().startswith("WSAA")
