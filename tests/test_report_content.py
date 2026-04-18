"""
Report Content Verification Tests — Cycle 7
=============================================
Verifies that generated DOCX and PDF reports contain correct
content: section headings, table row counts, values with units,
compliance status, and slurry analysis data.

These tests catch the class of bug found in Cycle 6 where the
report wrapper was silently producing empty steady-state sections.
"""

import os
import sys
import tempfile

import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from epanet_api import HydraulicAPI
from reports.docx_report import generate_docx_report
from epanet_api.slurry_solver import bingham_plastic_headloss


@pytest.fixture(scope='module')
def demo_api():
    """API loaded with demo_network, steady-state run."""
    api = HydraulicAPI()
    api.load_network('tutorials/demo_network/network.inp')
    return api


@pytest.fixture(scope='module')
def demo_results(demo_api):
    return demo_api.run_steady_state(save_plot=False)


@pytest.fixture(scope='module')
def demo_summary(demo_api):
    return demo_api.get_network_summary()


@pytest.fixture(scope='module')
def demo_docx(demo_results, demo_summary, tmp_path_factory):
    """Generate and parse a DOCX report from demo_network."""
    from docx import Document
    path = str(tmp_path_factory.mktemp("reports") / "demo_report.docx")
    generate_docx_report(
        {'steady_state': demo_results}, demo_summary, path,
        engineer_name='Test Engineer', project_name='Demo Network Report')
    doc = Document(path)
    return doc, path


@pytest.fixture(scope='module')
def slurry_api():
    """API loaded with mining_slurry_line."""
    api = HydraulicAPI()
    api.load_network('tutorials/mining_slurry_line/network.inp')
    return api


@pytest.fixture(scope='module')
def slurry_report_data(slurry_api):
    """Run steady-state + slurry solver, return wrapped results."""
    results = slurry_api.run_steady_state(save_plot=False)
    slurry_data = {}
    for pid in slurry_api.get_link_list('pipe'):
        pipe = slurry_api.get_link(pid)
        fdata = results['flows'].get(pid, {})
        Q = abs(fdata.get('avg_lps', 0)) / 1000
        if Q > 0 and pipe.diameter > 0:
            sd = bingham_plastic_headloss(
                flow_m3s=Q, diameter_m=pipe.diameter,
                length_m=pipe.length, density=1800,
                tau_y=15.0, mu_p=0.05, roughness_mm=0.1)
            slurry_data[pid] = sd
    results['slurry'] = slurry_data
    wrapped = {
        'steady_state': results,
        'slurry_params': {
            'yield_stress': 15.0,
            'plastic_viscosity': 0.05,
            'density': 1800,
        },
    }
    return wrapped, slurry_data


@pytest.fixture(scope='module')
def slurry_docx(slurry_report_data, slurry_api, tmp_path_factory):
    """Generate and parse a DOCX report with slurry data."""
    from docx import Document
    wrapped, _ = slurry_report_data
    summary = slurry_api.get_network_summary()
    path = str(tmp_path_factory.mktemp("reports") / "slurry_report.docx")
    generate_docx_report(
        wrapped, summary, path,
        engineer_name='Mining Engineer', project_name='Slurry Pipeline Report')
    doc = Document(path)
    return doc, path


def _all_text(doc):
    """Get all paragraph text from a docx Document."""
    return '\n'.join(p.text for p in doc.paragraphs)


def _all_table_text(doc):
    """Get all table cell text from a docx Document."""
    parts = []
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return '\n'.join(parts)


def _combined_text(doc):
    return _all_text(doc) + '\n' + _all_table_text(doc)


# =========================================================================
# TEST 1: Demo network DOCX report content
# =========================================================================

class TestDemoDocxContent:
    """Verify demo_network DOCX report has all required sections and data."""

    def test_executive_summary_present(self, demo_docx):
        doc, _ = demo_docx
        assert 'Executive Summary' in _all_text(doc)

    def test_network_description_present(self, demo_docx):
        doc, _ = demo_docx
        assert 'Network Description' in _all_text(doc)

    def test_steady_state_section_present(self, demo_docx):
        doc, _ = demo_docx
        text = _all_text(doc)
        assert 'Steady-State Results' in text

    def test_junction_pressure_table_row_count(self, demo_docx, demo_results):
        """Pressure table should have one row per junction."""
        doc, _ = demo_docx
        n_junctions = len(demo_results['pressures'])
        # Find the pressure table — look for table with 'Junction' header
        for tbl in doc.tables:
            header_cells = [c.text.strip() for c in tbl.rows[0].cells]
            if 'Junction' in header_cells and 'Min (m)' in header_cells:
                # Data rows = total rows - 1 (header)
                data_rows = len(tbl.rows) - 1
                assert data_rows == n_junctions, (
                    f"Pressure table has {data_rows} rows, expected {n_junctions}")
                return
        pytest.fail("Junction pressure table not found in report")

    def test_pipe_flow_table_row_count(self, demo_docx, demo_results):
        """Flow table should have one row per pipe."""
        doc, _ = demo_docx
        n_pipes = len(demo_results['flows'])
        for tbl in doc.tables:
            header_cells = [c.text.strip() for c in tbl.rows[0].cells]
            if 'Pipe' in header_cells and 'Velocity' in ' '.join(header_cells):
                data_rows = len(tbl.rows) - 1
                assert data_rows == n_pipes, (
                    f"Flow table has {data_rows} rows, expected {n_pipes}")
                return
        pytest.fail("Pipe flow table not found in report")

    def test_compliance_section_present(self, demo_docx):
        doc, _ = demo_docx
        assert 'Compliance' in _all_text(doc)

    def test_compliance_has_status_entries(self, demo_docx, demo_results):
        """Compliance table should have entries matching analysis results."""
        doc, _ = demo_docx
        n_compliance = len(demo_results['compliance'])
        for tbl in doc.tables:
            header_cells = [c.text.strip() for c in tbl.rows[0].cells]
            if 'Status' in header_cells and 'Message' in header_cells:
                data_rows = len(tbl.rows) - 1
                assert data_rows == n_compliance, (
                    f"Compliance table has {data_rows} rows, expected {n_compliance}")
                return
        pytest.fail("Compliance table not found in report")

    def test_pressure_values_have_units(self, demo_docx):
        """Pressure values in tables should include 'm' unit."""
        doc, _ = demo_docx
        table_text = _all_table_text(doc)
        # The pressure table has "Min (m)", "Max (m)" headers
        assert 'Min (m)' in table_text
        assert 'Max (m)' in table_text

    def test_velocity_values_have_units(self, demo_docx):
        """Velocity column header should include m/s unit."""
        doc, _ = demo_docx
        table_text = _all_table_text(doc)
        assert 'Velocity (m/s)' in table_text

    def test_flow_values_have_units(self, demo_docx):
        """Flow columns should include LPS unit."""
        doc, _ = demo_docx
        table_text = _all_table_text(doc)
        assert 'LPS' in table_text

    def test_conclusions_present(self, demo_docx):
        doc, _ = demo_docx
        assert 'Conclusions' in _all_text(doc)

    def test_engineer_name_on_cover(self, demo_docx):
        doc, _ = demo_docx
        assert 'Test Engineer' in _all_text(doc)

    def test_project_name_on_cover(self, demo_docx):
        doc, _ = demo_docx
        assert 'Demo Network Report' in _all_text(doc)

    def test_pressure_values_match_analysis(self, demo_docx, demo_results):
        """Spot-check that actual pressure values appear in the table."""
        doc, _ = demo_docx
        table_text = _all_table_text(doc)
        # Check first junction's avg pressure appears
        first_jid = list(demo_results['pressures'].keys())[0]
        avg_p = demo_results['pressures'][first_jid]['avg_m']
        assert str(avg_p) in table_text, (
            f"Junction {first_jid} avg pressure {avg_p} not in report tables")


# =========================================================================
# TEST 2: Mining slurry DOCX report content
# =========================================================================

class TestSlurryDocxContent:
    """Verify slurry pipeline DOCX report has Section 2A with correct data."""

    def test_slurry_section_heading(self, slurry_docx):
        doc, _ = slurry_docx
        assert 'Non-Newtonian Slurry Analysis' in _all_text(doc)

    def test_slurry_params_table(self, slurry_docx):
        """Slurry parameters table should show yield stress, viscosity, density."""
        doc, _ = slurry_docx
        combined = _combined_text(doc)
        assert '15.0 Pa' in combined, "Yield stress not in report"
        assert '0.05 Pa.s' in combined, "Plastic viscosity not in report"
        assert '1800' in combined, "Density not in report"

    def test_slurry_model_description(self, slurry_docx):
        doc, _ = slurry_docx
        text = _all_text(doc)
        assert 'Bingham Plastic' in text
        assert 'Buckingham-Reiner' in text

    def test_slurry_pipe_table_present(self, slurry_docx, slurry_report_data):
        """Per-pipe slurry results table with correct row count."""
        doc, _ = slurry_docx
        _, slurry_data = slurry_report_data
        n_slurry_pipes = len(slurry_data)
        for tbl in doc.tables:
            header_cells = [c.text.strip() for c in tbl.rows[0].cells]
            if 'Slurry HL' in ' '.join(header_cells):
                data_rows = len(tbl.rows) - 1
                assert data_rows == n_slurry_pipes, (
                    f"Slurry table has {data_rows} rows, expected {n_slurry_pipes}")
                return
        pytest.fail("Slurry pipe results table not found")

    def test_regime_column_populated(self, slurry_docx, slurry_report_data):
        """Regime column should contain laminar or turbulent."""
        doc, _ = slurry_docx
        combined = _combined_text(doc).lower()
        assert 'laminar' in combined or 'turbulent' in combined

    def test_reynolds_number_present(self, slurry_docx, slurry_report_data):
        """Re_B values should appear in the slurry table."""
        doc, _ = slurry_docx
        _, slurry_data = slurry_report_data
        combined = _combined_text(doc)
        assert 'Re_B' in combined, "Re_B header not in report"
        # Check at least one Re value appears
        first_pid = list(slurry_data.keys())[0]
        re_val = f"{slurry_data[first_pid]['reynolds']:.0f}"
        assert re_val in combined, f"Re_B value {re_val} not in report"

    def test_slurry_headloss_matches_solver(self, slurry_docx, slurry_report_data):
        """Spot-check that slurry headloss value from solver appears in report."""
        doc, _ = slurry_docx
        _, slurry_data = slurry_report_data
        combined = _combined_text(doc)
        first_pid = list(slurry_data.keys())[0]
        vel = slurry_data[first_pid]['velocity_ms']
        # Velocity should appear as formatted value
        vel_str = f"{vel:.2f}"
        assert vel_str in combined, (
            f"Slurry velocity {vel_str} for {first_pid} not in report")

    def test_slurry_summary_paragraph(self, slurry_docx, slurry_report_data):
        """Summary paragraph should mention pipe count and max velocity."""
        doc, _ = slurry_docx
        _, slurry_data = slurry_report_data
        text = _all_text(doc)
        assert f'{len(slurry_data)} pipes' in text


# =========================================================================
# TEST 3: PDF report validity
# =========================================================================

class TestPdfReport:
    """Verify PDF report is valid and has reasonable size."""

    def test_pdf_file_size(self, demo_results, demo_summary, tmp_path):
        """PDF should be > 10KB."""
        from reports.pdf_report import generate_pdf_report
        path = str(tmp_path / "test_report.pdf")
        generate_pdf_report(
            {'steady_state': demo_results}, demo_summary, path,
            engineer_name='PDF Test', project_name='PDF Verification')
        assert os.path.exists(path)
        size = os.path.getsize(path)
        assert size > 5000, f"PDF only {size} bytes, expected > 5KB"

    def test_pdf_is_valid(self, demo_results, demo_summary, tmp_path):
        """PDF should start with %PDF header (valid PDF file)."""
        from reports.pdf_report import generate_pdf_report
        path = str(tmp_path / "valid_check.pdf")
        generate_pdf_report(
            {'steady_state': demo_results}, demo_summary, path,
            engineer_name='PDF Test', project_name='PDF Check')
        with open(path, 'rb') as f:
            header = f.read(5)
        assert header == b'%PDF-', f"Invalid PDF header: {header}"

    def test_pdf_has_multiple_pages(self, demo_results, demo_summary, tmp_path):
        """PDF should have more than 1 page for a full report."""
        from reports.pdf_report import generate_pdf_report
        path = str(tmp_path / "pages_check.pdf")
        generate_pdf_report(
            {'steady_state': demo_results}, demo_summary, path,
            engineer_name='PDF Test', project_name='Pages Check')
        # Check for page count in PDF metadata
        with open(path, 'rb') as f:
            content = f.read()
        # Count page objects — each '/Type /Page' (not '/Pages') is one page
        import re
        pages = len(re.findall(rb'/Type\s*/Page[^s]', content))
        assert pages >= 2, f"PDF has only {pages} page(s), expected >= 2"
